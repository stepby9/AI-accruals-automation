import openai
import base64
import mimetypes
from pathlib import Path
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, asdict
from datetime import datetime, date
import json
import fitz  # PyMuPDF
from PIL import Image
import io

from config.settings import OpenAIConfig
from src.utils.logger import setup_logger
from src.utils.prompt_manager import get_system_prompt, get_user_prompt, get_model_config

logger = setup_logger(__name__)

@dataclass
class InvoiceData:
    bill_id: str
    invoice_number: Optional[str]
    invoice_date: Optional[date]  # Use date instead of datetime
    service_description: Optional[str]
    service_period: Optional[str]  # Text format like "January 2025", "Q1 2024"
    line_items_summary: Optional[str]  # Concatenated line items with amounts
    total_amount: Optional[float]
    currency: Optional[str]
    confidence_score: float
    extracted_at: datetime
    file_path: str

class InvoiceProcessor:
    def __init__(self):
        if not OpenAIConfig.API_KEY:
            raise ValueError("OpenAI API key not configured")
        
        # Create OpenAI client the correct way for v1.x
        from openai import OpenAI
        self.client = OpenAI(api_key=OpenAIConfig.API_KEY)
        self.model = OpenAIConfig.MODEL
        self.max_tokens = OpenAIConfig.MAX_TOKENS
        
        logger.info("Invoice processor initialized with OpenAI API")

    def process_invoice(self, file_path: str, bill_id: str) -> Optional[InvoiceData]:
        try:
            logger.info(f"Processing invoice: {file_path} for bill {bill_id}")
            
            file_type = self._get_file_type(file_path)
            
            if file_type == 'pdf':
                return self._process_pdf(file_path, bill_id)
            elif file_type in ['jpg', 'jpeg', 'png', 'gif', 'bmp']:
                return self._process_image(file_path, bill_id)
            elif file_type in ['xlsx', 'xls']:
                return self._process_excel(file_path, bill_id)
            elif file_type in ['docx', 'doc']:
                return self._process_word(file_path, bill_id)
            else:
                logger.warning(f"Unsupported file type: {file_type} for {file_path}")
                return None
                
        except Exception as e:
            logger.error(f"Error processing invoice {file_path}: {str(e)}")
            return None

    def _get_file_type(self, file_path: str) -> str:
        return Path(file_path).suffix.lower().lstrip('.')

    def _process_pdf(self, file_path: str, bill_id: str) -> Optional[InvoiceData]:
        try:
            doc = fitz.open(file_path)
            
            # Extract text from all pages
            text_content = ""
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text_content += page.get_text()
            
            # Convert first page to image for visual analysis
            first_page = doc.load_page(0)
            pix = first_page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom
            img_data = pix.tobytes("png")
            
            doc.close()
            
            return self._analyze_with_openai(
                text_content=text_content,
                image_data=img_data,
                file_path=file_path,
                bill_id=bill_id
            )
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}")
            return None

    def _process_image(self, file_path: str, bill_id: str) -> Optional[InvoiceData]:
        try:
            with open(file_path, 'rb') as f:
                image_data = f.read()
            
            return self._analyze_with_openai(
                text_content="",
                image_data=image_data,
                file_path=file_path,
                bill_id=bill_id
            )
            
        except Exception as e:
            logger.error(f"Error processing image {file_path}: {str(e)}")
            return None

    def _process_excel(self, file_path: str, bill_id: str) -> Optional[InvoiceData]:
        try:
            import pandas as pd
            
            # Read all sheets
            excel_data = pd.read_excel(file_path, sheet_name=None)
            
            text_content = "Excel file contents:\n"
            for sheet_name, df in excel_data.items():
                text_content += f"\n--- Sheet: {sheet_name} ---\n"
                text_content += df.to_string(max_rows=50)
            
            return self._analyze_with_openai(
                text_content=text_content,
                image_data=None,
                file_path=file_path,
                bill_id=bill_id
            )
            
        except Exception as e:
            logger.error(f"Error processing Excel {file_path}: {str(e)}")
            return None

    def _process_word(self, file_path: str, bill_id: str) -> Optional[InvoiceData]:
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text_content += row_text + "\n"
            
            return self._analyze_with_openai(
                text_content=text_content,
                image_data=None,
                file_path=file_path,
                bill_id=bill_id
            )
            
        except Exception as e:
            logger.error(f"Error processing Word document {file_path}: {str(e)}")
            return None

    def _analyze_with_openai(self, text_content: str, image_data: bytes, 
                           file_path: str, bill_id: str) -> Optional[InvoiceData]:
        try:
            # Prepare template variables
            template_vars = {'content_section': ''}
            
            # Get prompts and model config from prompt manager
            system_prompt = get_system_prompt("invoice_extraction")
            user_prompt = get_user_prompt("invoice_extraction", **template_vars)
            model_config = get_model_config("invoice_extraction")
            
            messages = [
                {
                    "role": "system",
                    "content": system_prompt
                },
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": user_prompt}
                    ]
                }
            ]
            
            if text_content:
                messages[1]["content"].append({
                    "type": "text", 
                    "text": f"\nExtracted text content:\n{text_content[:8000]}"  # Limit text length
                })
            
            if image_data:
                base64_image = base64.b64encode(image_data).decode('utf-8')
                messages[1]["content"].append({
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:image/png;base64,{base64_image}",
                        "detail": "high"
                    }
                })
            
            # Build API parameters dynamically based on model requirements
            api_params = {
                'model': model_config['model'],
                'messages': messages
            }
            
            # Add token limit parameter (varies by model)
            if 'max_completion_tokens' in model_config:
                api_params['max_completion_tokens'] = model_config['max_completion_tokens']
            elif 'max_tokens' in model_config:
                api_params['max_tokens'] = model_config['max_tokens']
            
            # Add temperature if supported
            if 'temperature' in model_config:
                api_params['temperature'] = model_config['temperature']
            
            response = self.client.chat.completions.create(**api_params)
            
            # Log token usage for performance monitoring
            if hasattr(response, 'usage') and response.usage:
                usage = response.usage
                logger.info(f"Token usage - Input: {usage.prompt_tokens}, Output: {usage.completion_tokens}, Total: {usage.total_tokens}")
            
            result = response.choices[0].message.content
            logger.debug(f"OpenAI response for {file_path}: {result}")
            
            # Clean the response - remove markdown code blocks if present
            cleaned_result = result.strip()
            if cleaned_result.startswith('```json'):
                cleaned_result = cleaned_result[7:]  # Remove ```json
            if cleaned_result.startswith('```'):
                cleaned_result = cleaned_result[3:]   # Remove ```
            if cleaned_result.endswith('```'):
                cleaned_result = cleaned_result[:-3]  # Remove closing ```
            cleaned_result = cleaned_result.strip()
            
            # Parse JSON response
            try:
                invoice_data_dict = json.loads(cleaned_result)
            except json.JSONDecodeError as json_err:
                logger.error(f"JSON parsing failed for {file_path}. Raw response: {result}")
                logger.error(f"Cleaned response: {cleaned_result}")
                logger.error(f"JSON error: {str(json_err)}")
                return None
            
            # Convert to InvoiceData object
            return self._dict_to_invoice_data(invoice_data_dict, bill_id, file_path)
            
        except Exception as e:
            logger.error(f"Error analyzing with OpenAI for {file_path}: {str(e)}")
            return None


    def _dict_to_invoice_data(self, data_dict: Dict, bill_id: str, file_path: str) -> InvoiceData:
        # Create line items summary
        line_items = data_dict.get('line_items', [])
        line_items_summary = None
        if line_items:
            items_text = []
            for item in line_items:
                desc = item.get('description', 'N/A')
                amount = item.get('amount', 'N/A')
                items_text.append(f"{desc}: {amount}")
            line_items_summary = "; ".join(items_text)
        
        return InvoiceData(
            bill_id=bill_id,
            invoice_number=data_dict.get('invoice_number'),
            invoice_date=self._parse_date(data_dict.get('invoice_date')),
            service_description=data_dict.get('service_description'),
            service_period=data_dict.get('service_period'),  # Keep as text
            line_items_summary=line_items_summary,
            total_amount=float(data_dict.get('total_amount')) if data_dict.get('total_amount') is not None else None,
            currency=data_dict.get('currency'),
            confidence_score=float(data_dict.get('confidence_score', 0.5)),
            extracted_at=datetime.now(),
            file_path=file_path
        )

    def _parse_date(self, date_str: str) -> Optional[date]:
        if not date_str:
            return None
        try:
            return datetime.strptime(date_str, '%Y-%m-%d').date()
        except (ValueError, TypeError):
            return None

    def process_multiple_invoices(self, invoice_files: List[str], bill_id: str) -> List[InvoiceData]:
        results = []
        
        for file_path in invoice_files:
            result = self.process_invoice(file_path, bill_id)
            if result:
                results.append(result)
        
        logger.info(f"Processed {len(results)} invoices successfully for bill {bill_id}")
        return results

    def is_invoice_already_processed(self, bill_id: str, database) -> bool:
        """Check if invoice for this bill is already processed"""
        try:
            # This will be implemented when database is created
            # For now, return False to process all invoices
            return False
        except Exception as e:
            logger.error(f"Error checking if invoice processed for bill {bill_id}: {str(e)}")
            return False